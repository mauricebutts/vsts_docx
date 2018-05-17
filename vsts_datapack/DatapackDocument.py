""" Author: Maurice Butts
Date: 5/16/2018
Desc: This file acts as the API to the Datapack package. Use this to interface with the custom built commands.
"""
from vsts_datapack.DatapackDocx.DatapackDocx import _create_docx_table_from_query
from vsts_datapack.DatapackDocx.DatapackDocx import _add_datapack_hyperlink
from vsts_datapack.DatapackVsts.DatapackVsts import _datapack_query
from vsts_datapack.DatapackVsts.DatapackVsts import _datapack_item_count_query
import docx
from docx.shared import Inches
import json

# TODO: Use a docx template to read from?
# TODO: Sort the tables


class DatapackDocument:

    def __init__(self, config_path=None):
        self.__create()

        # If no config is declared let's check current dir, if no config create a new one
        if not config_path:
            try:
                with open('config.json') as config_file:
                    self.config = json.load(config_file)
            except Exception:
                print('No config found, creating config in current dir')
                print('Please fill in the config file with your vsts creds')
                self.__create_config()
        else:
            print('Grabbing config from {}'.format(config_path))
            self.__grab_config(config_path)

    def __create(self):
        self.document = docx.Document()
        self.paragraph_header = self.document

    def __grab_config(self, config_path):
        with open(config_path, 'w') as read_config:
            self.config = json.load(read_config)

    def __create_config(self):
        config_format = {
           1: ['token', 'team_instance', 'project']
        }

        with open('config.json', 'w') as write_config:
            json.dump(config_format, write_config)

    def add_paragraph(self, text):
        self.__add_new_paragraph_and_set_header(text)

    def add_text(self, text):
        if self.paragraph_header == self.document:
            raise Exception("No text based object at head. Please start a paragraph or something.")
        self.paragraph_header.add_run(text)

    def add_linebreak(self):
        if self.paragraph_header == self.document:
            raise Exception("No text started, please add a paragraph.")
        self.__add_new_paragraph_and_set_header('')

    def page_break(self):
        self.document.add_page_break()
        self.__add_new_paragraph_and_set_header('')

    def write(self, path):
        self.document.save(path)

    def add_image(self, path, image_width):
        self.document.add_picture(path, width=Inches(image_width))

    def add_heading(self, text, heading_type):
        self.document.add_heading(text, heading_type)
        self.__add_new_paragraph_and_set_header('')

    def create_table_from_query(self, vsts_config_key, query_id):
        data_list_of_lists = self._vsts_query(vsts_config_key, query_id)
        column_names = ['work item ids', 'returned titles', 'returned priority', 'returned state', 'returned tag']
        self.create_docx_table(list_of_column_names=column_names, list_of_data_lists=data_list_of_lists)

    def add_workitem_count_hyperlink(self, url_text, vsts_config_key, query_id):
        vsts_creds = self.__grab_vsts_creds(str(vsts_config_key))
        return _add_datapack_hyperlink(paragraph=self.paragraph_header, url_text=url_text, query_id=query_id,
                                       token=vsts_creds[0], team_instance=vsts_creds[1], project=vsts_creds[2])

    def _vsts_query(self, vsts_config_key, query_id):
        vsts_creds = self.__grab_vsts_creds(str(vsts_config_key))
        return _datapack_query(vsts_creds[0], vsts_creds[1], query_id)

    def _vsts_item_count_query(self, vsts_config_key, query_id):
        vsts_creds = self.__grab_vsts_creds(str(vsts_config_key))
        return _datapack_item_count_query(vsts_creds[0], vsts_creds[1], query_id, vsts_creds[2])

    def create_docx_table(self, list_of_column_names, list_of_data_lists):
        return _create_docx_table_from_query(self.document, list_of_column_names, list_of_data_lists)

    def __grab_vsts_creds(self, config_key):
        return self.config[config_key]

    def __add_new_paragraph_and_set_header(self, text):
        paragraph = self.document.add_paragraph(text)
        self.paragraph_header = paragraph
