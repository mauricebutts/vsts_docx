from DatapackVsts import DatapackVsts as av
from docx.enum.dml import MSO_THEME_COLOR_INDEX
import docx


def _create_docx_table_from_query(document, list_of_column_names, list_of_data_lists):
    """
    :param document             : class         : Document object to be acted on.
    :param list_of_column_names : list          : WorkItem Ids
    :param list_of_data_lists   : list of lists : N number of lists to be turned into columns and created into the DatapackDocx table
    :return:                      class         : DatapackDocx.document.add_table
    """

    # First need to find the maximum number of rows needed. Will look for the largest list in list_of_data_lists,
    # although we should expect to see the same length between all lists.
    largest_list = -1
    for list in list_of_data_lists:
        if len(list) > largest_list:
            largest_list = len(list)

    number_of_rows = largest_list

    # Construct table object. Number_of_rows + 1 to allow for column headers.
    docx_table = document.add_table(rows=number_of_rows+1, cols=len(list_of_data_lists))

    # First create column names
    for idx, name in enumerate(list_of_column_names):
        cell = docx_table.cell(0, idx)
        cell.text = name

    # Load data into table. Converting everything to str() in order to conform to docx_table.cell() formatting.
    column_idx = 0
    for list in list_of_data_lists:
        for idx, item in enumerate(list):
            cell = docx_table.cell(idx+1, column_idx)
            cell.text = str(item)
        column_idx += 1

    docx_table.style = 'Table Grid'


def _add_hyperlink(paragraph, text, url):

    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element and a new w:rPr element
    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    # Create a new Run object and add the hyperlink into it
    r = paragraph.add_run ()
    r._r.append (hyperlink)

    # A workaround for the lack of a hyperlink style (doesn't go purple after using the link)
    # Delete this if using a template that has the hyperlink style in it
    r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
    r.font.underline = True

    return hyperlink


def _add_datapack_hyperlink(paragraph, url_text, query_id, token, team_instance, project):

    query_count, url = av.datapack_item_count_query(token, team_instance, query_id, project)
    query_count = str(query_count)

    text = query_count + ' ' + url_text

    _add_hyperlink(paragraph, text, url)
